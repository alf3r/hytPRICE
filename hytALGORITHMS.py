import pathlib
import os
import pandas as pd
import numpy as np
import time

import openpyxl.utils
from openpyxl.worksheet.cell_range import CellRange
from openpyxl.styles import NamedStyle
from openpyxl.styles import Font
from openpyxl.utils import get_column_letter

from docxtpl import DocxTemplate, InlineImage

########################################################################################################################
# Удаление форматов из листа
def removeFormatting(ws):
    # ws is not the worksheet name, but the worksheet object
    for row in ws.iter_rows():
        for cell in row:
            cell.style = 'Normal'

########################################################################################################################
# Сравнение двух pandas датафреймов
def report_diff(x):
    """Function to use with groupby.apply to highlight value changes."""
    return x.iloc[0] if x.iloc[0] == x.iloc[1] or pd.isna(x).all() else f'{x.iloc[0]} ---> {x.iloc[1]}'

def strip(x):
    """Function to use with applymap to strip whitespaces from a dataframe."""
    return x.strip() if isinstance(x, str) else x

def diff_pd(old_df, new_df, idx_col):
    """
    Identify differences between two pandas DataFrames using a key column.

    Key column is assumed to have a unique row identifier, i.e. no duplicates.

    Args:
        old_df (pd.DataFrame): first dataframe
        new_df (pd.DataFrame): second dataframe
        idx_col (str|list(str)): column name(s) of the index,
          needs to be present in both DataFrames
    Example:
        diff_pd(df_1, df_2, ["ID"]
    """
    # setting the column name as index for fast operations
    old_df = old_df.set_index(idx_col)
    new_df = new_df.set_index(idx_col)
    # get the added and removed rows
    old_keys = old_df.index
    new_keys = new_df.index
    if isinstance(old_keys, pd.MultiIndex):
        removed_keys = old_keys.difference(new_keys)
        added_keys = new_keys.difference(old_keys)
    else:
        removed_keys = np.setdiff1d(old_keys, new_keys)
        added_keys = np.setdiff1d(new_keys, old_keys)
    # populate the output data with non empty dataframes
    out_data = {}
    removed = old_df.loc[removed_keys]
    if not removed.empty:
        out_data["removed"] = removed
    added = new_df.loc[added_keys]
    if not added.empty:
        out_data["added"] = added
    # focusing on common data of both dataframes
    common_keys = np.intersect1d(old_keys, new_keys, assume_unique=True)
    common_columns = np.intersect1d(
        old_df.columns, new_df.columns, assume_unique=True
    )
    new_common = new_df.loc[common_keys, common_columns].map(strip)
    old_common = old_df.loc[common_keys, common_columns].map(strip)
    # get the changed rows keys by dropping identical rows
    # (indexes are ignored, so we'll reset them)
    common_data = pd.concat(
        [old_common.reset_index(), new_common.reset_index()], sort=True
    )
    changed_keys = common_data.drop_duplicates(keep=False)[idx_col]
    if isinstance(changed_keys, pd.Series):
        changed_keys = changed_keys.unique()
    else:
        changed_keys = changed_keys.drop_duplicates().set_index(idx_col).index
    # combining the changed rows via multi level columns
    df_all_changes = pd.concat(
        [old_common.loc[changed_keys], new_common.loc[changed_keys]],
        axis='columns',
        keys=['old', 'new']
    ).swaplevel(axis='columns')
    # using report_diff to merge the changes in a single cell with "-->"
    df_changed = df_all_changes.groupby(level=0, axis=1).apply(lambda frame: frame.apply(report_diff, axis=1))
    df_notChanged = new_common
    # add changed dataframe to output data only if non empty
    if not df_changed.empty:
        out_data['changed'] = df_changed
    if not df_notChanged.empty:
        out_data['notChanged'] = df_notChanged
    combined_df = pd.DataFrame()
    for sname, data in out_data.items():
        data.loc[data['MODEL RUS'] != '', 'STATUS'] = sname
        combined_df = pd.concat([data, combined_df])
    combined_df = combined_df[['STATUS', 'MODEL RUS', 'DESCRIPTION ENG', 'DESCRIPTION RUS',
                           'GPL, CNY', 'РРЦ, CNY', 'COMMENTS', 'MAINTYPE', 'SUBTYPE', 'FREQUENCY BAND']
    ]
    combined_df = combined_df[~combined_df.index.duplicated(keep='last')]
    combined_df = combined_df.sort_values(by="DESCRIPTION RUS")
    return combined_df


#######################################################################################################################
# Сравнение прайсов RUS
def compareEXLSrus(NewFileName: str, OldFileName: str):
    errors = []
    from openpyxl.styles import PatternFill
    wbNEW = openpyxl.load_workbook(NewFileName, data_only=True)
    mySheets = wbNEW.sheetnames
    wbNEW.close()
    idxXLSX = NewFileName.find('.xlsx')
    ComparedFileName = NewFileName[:idxXLSX] + '-compared' + NewFileName[idxXLSX:]
    wbCOMPARED = openpyxl.Workbook()
    wbCOMPARED.save(filename=ComparedFileName)
    wbCOMPARED.close()
    for n, item1 in enumerate(mySheets):
        try:
            sheet_2_compare = item1
            new_df = pd.read_excel(NewFileName, sheet_2_compare)
            old_df = pd.read_excel(OldFileName, sheet_2_compare)
            diff_df = diff_pd(old_df, new_df, 'BOM')
            with pd.ExcelWriter(ComparedFileName, engine='openpyxl', mode='a') as writer:
                diff_df.to_excel(writer, index=True, sheet_name=sheet_2_compare)
        except Exception as myError:
            errors.append('Ошибка: ' + str(myError))
            print('Ошибка: ' + str(myError))
    # раскрасить ячейки
    redFill = PatternFill(start_color='ff0000',
                          end_color='ff0000',
                          fill_type='solid')
    greenFill = PatternFill(start_color='7FFF00',
                          end_color='7FFF00',
                          fill_type='solid')
    yellowFill = PatternFill(start_color='FFD700',
                          end_color='FFD700',
                          fill_type='solid')
    grayFill = PatternFill(start_color='B8B8B8',
                          end_color='B8B8B8',
                          fill_type='solid')
    wbCOMPARED = openpyxl.load_workbook(ComparedFileName, data_only=True)
    del wbCOMPARED['Sheet']
    compSheets = wbCOMPARED.sheetnames
    for n, item1 in enumerate(compSheets):
        sheetIn = wbCOMPARED[item1]
        rowsIn = sheetIn.max_row+1
        for row1 in range(2, rowsIn):
            rowFlag = sheetIn.cell(row=row1, column=2).value
            if (rowFlag == 'removed'):
                sheetIn.cell(row=row1, column=2).fill = redFill
            elif (rowFlag == 'changed'):
                sheetIn.cell(row=row1, column=2).fill = yellowFill
            elif (rowFlag == 'added'):
                sheetIn.cell(row=row1, column=2).fill = greenFill
            elif (rowFlag == 'notChanged'):
                sheetIn.cell(row=row1, column=2).fill = grayFill
            else:
                errors.append('Ошибка ->> нет такого статуса ' + rowFlag)
                print('Ошибка ->> нет такого статуса ' + rowFlag)
    wbCOMPARED.save(filename=ComparedFileName)
    wbCOMPARED.close()
    print('Успех! ' + compareEXLSrus.__name__)
    return errors


########################################################################################################################
# Комбинирование нескольких ТЕРМИНАЛЛЬНЫХ прайсов в один ТЕРМИНАЛЬНЫЙ ПРАЙС с закладками
def cpqPricesCombine(filesPrices: list[str], wbCombine: openpyxl.Workbook):
    errors = []
    for n, item in enumerate(filesPrices):
        wbIn = openpyxl.load_workbook(item, data_only=True)
        sheetInName = wbIn.sheetnames[0]
        sheetInName = sheetInName.replace("\xa0", "")
        sheetInName = sheetInName.replace(" ", "")
        sheetInName = sheetInName.replace("-", "")
        sheetIn = wbIn.active
        sheetOut = wbCombine.create_sheet(sheetInName)
        # подготовим листы для соединения
        errors.append(prepare_sheet(sheetIn, sheetOut, sheetInName))
        wbIn.close()
    print('Успех! ' + cpqPricesCombine.__name__)
    return errors

def copy_sheet(source_sheet, target_sheet):
    copy_cells(source_sheet, target_sheet)  # copy all the cel values and styles
    copy_sheet_attributes(source_sheet, target_sheet)

def copy_sheet_attributes(source_sheet, target_sheet):
    from copy import copy
    target_sheet.sheet_format = copy(source_sheet.sheet_format)
    target_sheet.sheet_properties = copy(source_sheet.sheet_properties)
    target_sheet.merged_cells = copy(source_sheet.merged_cells)
    target_sheet.page_margins = copy(source_sheet.page_margins)
    target_sheet.freeze_panes = copy(source_sheet.freeze_panes)

    # set row dimensions
    # So you cannot copy the row_dimensions attribute. Does not work (because of meta data in the attribute I think). So we copy every row's row_dimensions. That seems to work.
    for rn in range(len(source_sheet.row_dimensions)):
        target_sheet.row_dimensions[rn] = copy(source_sheet.row_dimensions[rn])

    if source_sheet.sheet_format.defaultColWidth is None:
        print('Unable to copy default column wide')
    else:
        target_sheet.sheet_format.defaultColWidth = copy(source_sheet.sheet_format.defaultColWidth)

    # set specific column width and hidden property
    # we cannot copy the entire column_dimensions attribute so we copy selected attributes
    for key, value in source_sheet.column_dimensions.items():
        target_sheet.column_dimensions[key].min = copy(source_sheet.column_dimensions[key].min)   # Excel actually groups multiple columns under 1 key. Use the min max attribute to also group the columns in the targetSheet
        target_sheet.column_dimensions[key].max = copy(source_sheet.column_dimensions[key].max)  # https://stackoverflow.com/questions/36417278/openpyxl-can-not-read-consecutive-hidden-columns discussed the issue. Note that this is also the case for the width, not onl;y the hidden property
        target_sheet.column_dimensions[key].width = copy(source_sheet.column_dimensions[key].width) # set width for every column
        target_sheet.column_dimensions[key].hidden = copy(source_sheet.column_dimensions[key].hidden)

def copy_cells(source_sheet, target_sheet):
    from copy import copy
    for (row, col), source_cell in source_sheet._cells.items():
        target_cell = target_sheet.cell(column=col, row=row)

        target_cell._value = source_cell._value
        target_cell.data_type = source_cell.data_type

        if source_cell.has_style:
            target_cell.font = copy(source_cell.font)
            target_cell.border = copy(source_cell.border)
            target_cell.fill = copy(source_cell.fill)
            target_cell.number_format = copy(source_cell.number_format)
            target_cell.protection = copy(source_cell.protection)
            target_cell.alignment = copy(source_cell.alignment)

        if source_cell.hyperlink:
            target_cell._hyperlink = copy(source_cell.hyperlink)

        if source_cell.comment:
            target_cell.comment = copy(source_cell.comment)

########################################################################################################################
# Комбинирование нескольких СИСТЕМНЫХ прайсов в один СИСТЕМНЫЙ прайс с закладками
def cpqSystemCombine(filesPrices: list[str], wb_target: openpyxl.Workbook):
    for n, item in enumerate(filesPrices):
        wb_source = openpyxl.load_workbook(item, data_only = True)
        source_sheet = wb_source.active
        sheetInName = wb_source.sheetnames[0]
        sheetInName = sheetInName.replace("\xa0", "")
        sheetInName = sheetInName.replace(" ", "")
        sheetInName = sheetInName.replace("-", "")
        target_sheet = wb_target.create_sheet(sheetInName)
        copy_sheet(source_sheet, target_sheet)
        if 'Sheet' in wb_target.sheetnames:  # remove default sheet
            wb_target.remove(wb_target['Sheet'])
        wb_source.close()
    print('Успех! ' + cpqSystemCombine.__name__)


########################################################################################################################
# подготовка СИСТЕМНОГО листа для прайса
def prepare_SYSTEMsheet(sheetIn, sheetOut):
    currColA = openpyxl.utils.column_index_from_string('A')
    currColB = openpyxl.utils.column_index_from_string('B')
    currColC = openpyxl.utils.column_index_from_string('C')
    currColD = openpyxl.utils.column_index_from_string('D')
    currColE = openpyxl.utils.column_index_from_string('E')
    currColF = openpyxl.utils.column_index_from_string('F')
    currColG = openpyxl.utils.column_index_from_string('G')
    rowsIn = sheetIn.max_row + 1
    # копируем данные из столбца 9 в столбец 1 в новом прайсе
    currRow = 2
    for row1 in range(2, rowsIn):
        rowFlag = sheetIn.cell(row=row1, column=12).value
        if (rowFlag == '可选项') or (rowFlag == ''):
            continue
        valueIn = sheetIn.cell(row=row1, column=9).value
        if (valueIn == 'SelfdevelopedSoftware')or (valueIn == 'SpecializedPurchasedSoftware') or (valueIn == 'GeneralPurchasedSoftware'):
            valueIn = 'Software'
        elif (valueIn == 'SelfdevelopedHardware') or (valueIn == 'SpecializedPurchasedHardware') or (valueIn == 'GeneralPurchasedHardware') or (valueIn == ''):
            valueIn = 'Hardware'
        cellOut = sheetOut.cell(row=currRow, column=1)
        sheetOut[cellOut.coordinate].value = valueIn

        valueIn = sheetIn.cell(row=row1, column=2).value
        cellOut = sheetOut.cell(row=currRow, column=2)
        sheetOut[cellOut.coordinate].value = valueIn

        valueIn = sheetIn.cell(row=row1, column=3).value
        cellOut = sheetOut.cell(row=currRow, column=3)
        sheetOut[cellOut.coordinate].value = valueIn

        valueIn = sheetIn.cell(row=row1, column=4).value
        cellOut = sheetOut.cell(row=currRow, column=4)
        sheetOut[cellOut.coordinate].value = valueIn

        valueIn = sheetIn.cell(row=row1, column=7).value
        cellOut = sheetOut.cell(row=currRow, column=7)
        sheetOut[cellOut.coordinate].value = valueIn
        currRow = currRow + 1

    # вставим заглавную строку
    sheetOut.cell(row=1, column=currColA).value = 'type'
    sheetOut.cell(row=1, column=currColB).value = 'itemCode'
    sheetOut.cell(row=1, column=currColC).value = 'itemName'
    sheetOut.cell(row=1, column=currColD).value = 'itemDescription'
    sheetOut.cell(row=1, column=currColE).value = 'productModel'
    sheetOut.cell(row=1, column=currColF).value = 'itemNum'
    sheetOut.cell(row=1, column=currColG).value = 'price'
    #print('Успех! ' + prepare_SYSTEMsheet.__name__)


########################################################################################################################
# подготовка СИСТЕМНОГО листа для прайса
def sheetCPQ2RUS(sheetIn, sheetOut, currentCurrency):
    from openpyxl.styles import PatternFill
    from openpyxl.styles import Font
    from openpyxl.styles import Alignment
    yellowFill = PatternFill(start_color='FFD700',
                          end_color='FFD700',
                          fill_type='solid')
    grayFill = PatternFill(start_color='B8B8B8',
                          end_color='B8B8B8',
                          fill_type='solid')
    bold_font = Font(bold=True)
    center_align = Alignment(
        horizontal="center",
        vertical="center"
    )
    right_align = Alignment(
        horizontal="right",
        vertical="center"
    )
    currentCurrency = currentCurrency.replace(",", ".")
    currColA = openpyxl.utils.column_index_from_string('A') #1
    currColB = openpyxl.utils.column_index_from_string('B') #2
    currColC = openpyxl.utils.column_index_from_string('C') #3
    currColD = openpyxl.utils.column_index_from_string('D') #4
    currColE = openpyxl.utils.column_index_from_string('E') #5
    currColF = openpyxl.utils.column_index_from_string('F') #6
    currColG = openpyxl.utils.column_index_from_string('G') #7
    currColH = openpyxl.utils.column_index_from_string('H') #8
    currColI = openpyxl.utils.column_index_from_string('I') #9
    currColJ = openpyxl.utils.column_index_from_string('J') #10
    currColK = openpyxl.utils.column_index_from_string('K') #11
    currColL = openpyxl.utils.column_index_from_string('L') #12
    currColM = openpyxl.utils.column_index_from_string('M') #13
    rowsIn = sheetIn.max_row + 1
    # копируем данные из столбца 9 в столбец 1 в новом прайсе
    for col_range in range(1, 7):
        cell_title = sheetOut.cell(1, col_range)
        cell_title.fill = grayFill
        cell_title.font = bold_font
        cell_title.alignment = center_align
    currRow = 2
    for row1 in range(2, rowsIn):
        rowFlag = sheetIn.cell(row=row1, column=currColL).value
        if (rowFlag == '可选项') or (rowFlag == ''):
            for col_range in range(1, 7):
                cell_part = sheetOut.cell(row1, col_range)
                cell_part.fill = yellowFill
                cell_part.font = bold_font

        # itemType = sheetIn.cell(row=row1, column=currColI).value
        # if (itemType == 'SelfdevelopedSoftware')or (itemType == 'SpecializedPurchasedSoftware') or (itemType == 'GeneralPurchasedSoftware'):
        #     itemType = 'Software'
        # elif (itemType == 'SelfdevelopedHardware') or (itemType == 'SpecializedPurchasedHardware') or (itemType == 'GeneralPurchasedHardware') or (itemType == ''):
        #     itemType = 'Hardware'
        # cellOut = sheetOut.cell(row=currRow, column=currColA)
        # sheetOut[cellOut.coordinate].value = itemType

        # перенесем BOM из столбца "B" в столбец "C"
        valueIn = sheetIn.cell(row=row1, column=currColB).value
        cellOut = sheetOut.cell(row=currRow, column=currColC)
        cellOut.alignment = right_align
        sheetOut[cellOut.coordinate].value = valueIn

        # перенесем itemName из столбца "C" в столбец "B"
        valueIn = sheetIn.cell(row=row1, column=currColC).value
        cellOut = sheetOut.cell(row=currRow, column=currColB)
        sheetOut[cellOut.coordinate].value = valueIn

        # перенесем Description из столбца "D" в столбец "A"
        description = sheetIn.cell(row=row1, column=currColD).value
        cellOut = sheetOut.cell(row=currRow, column=currColA)
        description = description.replace("~", "...")
        description = description.replace("～", "...")
        sheetOut[cellOut.coordinate].value = description

        # # перенесем Quontity из столбца "E" в столбец "C"
        # valueIn = sheetIn.cell(row=row1, column=currColE).value
        # cellOut = sheetOut.cell(row=currRow, column=currColC)
        # sheetOut[cellOut.coordinate].value = valueIn

        # # перенесем Unit из столбца "F" в столбец "D"
        # valueIn = sheetIn.cell(row=row1, column=currColF).value
        # cellOut = sheetOut.cell(row=currRow, column=currColD)
        # sheetOut[cellOut.coordinate].value = valueIn

        # перенесем Currency из столбца "H" в столбец "E"
        currency = sheetIn.cell(row=row1, column=currColH).value
        cellOut = sheetOut.cell(row=currRow, column=currColE)
        cellOut.alignment = center_align
        if currency == 'USD':
            sheetOut[cellOut.coordinate].value = 'CNY'
        else:
            if (rowFlag == '可选项') or (rowFlag == ''):
                sheetOut[cellOut.coordinate].value = ''
            else:
                sheetOut[cellOut.coordinate].value = 'unknown'

        # перенесем RRP из столбца "G" в столбец "D" и умножим на курс
        pricePPR = sheetIn.cell(row=row1, column=currColG).value
        cellOut = sheetOut.cell(row=currRow, column=currColD)
        cellOut.alignment = right_align
        if currency == 'USD':
            pricePPR = pricePPR.replace(".00000", "")
            priceCNY = (round(float(pricePPR) * float(currentCurrency), 2))
            # priceCNY = priceCNY.replace(".", ",")
            sheetOut[cellOut.coordinate].value = priceCNY
        else:
            sheetOut[cellOut.coordinate].value = pricePPR

        # # Скопируем столбец "А"
        # valueIn = sheetIn.cell(row=row1, column=currColA).value
        # cellOut = sheetOut.cell(row=currRow, column=currColA)
        # sheetOut[cellOut.coordinate].value = valueIn

        # перенесем Item Property из столбца "I" в столбец "F" и умножим на курс
        valueIn = sheetIn.cell(row=row1, column=currColI).value
        cellOut = sheetOut.cell(row=currRow, column=currColF)
        sheetOut[cellOut.coordinate].value = valueIn

        # # Скопируем столбец "J"
        # priceTotal = sheetIn.cell(row=row1, column=currColJ).value
        # cellOut = sheetOut.cell(row=currRow, column=currColJ)
        # if currency == 'USD':
        #     if priceTotal == '':
        #         priceTotal = 0
        #     else:
        #         priceTotal = priceTotal.replace(".00000", "")
        #     sheetOut[cellOut.coordinate].value = str(round(int(priceTotal) * int(currentCurrency), 2))
        # else:
        #     sheetOut[cellOut.coordinate].value = priceTotal

        # # Скопируем столбец "K"
        # valueIn = sheetIn.cell(row=row1, column=currColK).value
        # cellOut = sheetOut.cell(row=currRow, column=currColK)
        # sheetOut[cellOut.coordinate].value = valueIn
        #
        # # Скопируем столбец "L"
        # valueIn = sheetIn.cell(row=row1, column=currColL).value
        # cellOut = sheetOut.cell(row=currRow, column=currColL)
        # sheetOut[cellOut.coordinate].value = valueIn
        #
        # # Скопируем столбец "M"
        # valueIn = sheetIn.cell(row=row1, column=currColM).value
        # cellOut = sheetOut.cell(row=currRow, column=currColM)
        # sheetOut[cellOut.coordinate].value = valueIn

        sheetOut.column_dimensions['A'].width = 80
        sheetOut.column_dimensions['B'].width = 20
        sheetOut.column_dimensions['C'].width = 17
        sheetOut.column_dimensions['D'].width = 11
        sheetOut.column_dimensions['E'].width = 9
        sheetOut.column_dimensions['F'].width = 30
        currRow = currRow + 1

    # вставим заглавную строку
    sheetOut.cell(row=1, column=currColA).value = 'description'
    sheetOut.cell(row=1, column=currColB).value = 'itemName'
    sheetOut.cell(row=1, column=currColC).value = 'itemCode *'
    sheetOut.cell(row=1, column=currColD).value = 'RRP'
    sheetOut.cell(row=1, column=currColE).value = 'Currency'
    sheetOut.cell(row=1, column=currColF).value = 'Item property'
    # sheetOut.cell(row=1, column=currColA).value = 'Serial number'
    # sheetOut.cell(row=1, column=currColC).value = 'Quantity *'
    # sheetOut.cell(row=1, column=currColD).value = 'Unit'
    # sheetOut.cell(row=1, column=currColJ).value = 'Row total'
    # sheetOut.cell(row=1, column=currColK).value = 'Associated software'
    # sheetOut.cell(row=1, column=currColL).value = 'Flag'
    # sheetOut.cell(row=1, column=currColM).value = 'Configuration rules'
    #print('Успех! ' + prepare_SYSTEMsheet.__name__)


########################################################################################################################
# подготовка листа для прайса
def prepare_sheet(sheetIn, sheetOut, sheetInName):
    import openpyxl.cell as Cell
    errors = []
    currColA = openpyxl.utils.column_index_from_string('A')
    currColB = openpyxl.utils.column_index_from_string('B')
    currColC = openpyxl.utils.column_index_from_string('C')
    currColD = openpyxl.utils.column_index_from_string('D')
    currColE = openpyxl.utils.column_index_from_string('E')
    currColF = openpyxl.utils.column_index_from_string('F')
    currColG = openpyxl.utils.column_index_from_string('G')
    currColH = openpyxl.utils.column_index_from_string('H')
    rowsIn = sheetIn.max_row + 1
    colsIn = openpyxl.utils.column_index_from_string('G') + 1
    # разъединяем соединенные ячейки
    all_merged_cell_ranges: list[CellRange] = list(
        sheetIn.merged_cells.ranges
    )
    # во все разъединенные на предыдущем шаге копируем значения, которые были в объединенной ячейке
    for merged_cell_range in all_merged_cell_ranges:
        merged_cell: Cell = merged_cell_range.start_cell
        sheetIn.unmerge_cells(range_string=merged_cell_range.coord)
        # Don't need to convert iterator to list here since `merged_cell_range` is cached
        for row_index, col_index in merged_cell_range.cells:
            cell: Cell = sheetIn.cell(row=row_index, column=col_index)
            cell.value = merged_cell.value
    # копируем данные из листа CPQ в новый прайс - без первой строки
    for row1 in range(2, rowsIn):
        for col1 in range(1, colsIn):
            valueIn = sheetIn.cell(row=row1, column=col1).value
            cellOut = sheetOut.cell(row=row1 - 1, column=col1)
            sheetOut[cellOut.coordinate].value = valueIn
    # создаем новый столбец А (МОДЕЛЬ)
    sheetOut.insert_cols(0)
    # переносим данные из столбца F(productModel) в новый столбец A(МОДЕЛЬ), если там нет данных, то берем из G(
    # itemNum) или из C(itemCode)
    for row1 in range(1, sheetIn.max_row):
        currValue = sheetOut.cell(row=row1, column=currColF).value
        if (currValue is None) or (currValue == ''):
            currValue = sheetOut.cell(row=row1, column=currColG).value
            if (currValue is None) or (currValue == ''):
                currValue = sheetOut.cell(row=row1, column=currColC).value
        sheetOut.cell(row=row1, column=currColA).value = currValue
    # удаляем нули из цен в столбце H
    for row1 in range(1, sheetIn.max_row):
        currValue = sheetOut.cell(row=row1, column=currColH).value
        try:
            newValue = currValue.split('.')[0]
        except:
            newValue = 0
        if newValue == 'null':
            newValue = 0
        if newValue == '0E-8':
            newValue = 0
        try:
            sheetOut.cell(row=row1, column=currColH).value = int(newValue)
        except:
            print(newValue)
            continue
        sheetOut.cell(row=row1, column=currColH).number_format = '[$$-409]#,##0.00;[RED]-[$$-409]#,##0.00'
    # исправляем названия в столбце A (модель)
    rowAnt = 1
    for row1 in range(1, sheetIn.max_row):
        colAvalue = str(sheetOut.cell(row=row1, column=currColA).value)
        colCvalue = str(sheetOut.cell(row=row1, column=currColC).value)
        colEvalue = str(sheetOut.cell(row=row1, column=currColE).value)
        colType = find_item_type(colCvalue, colAvalue, colEvalue)
        if colType == 1:  # терминалы
            model = colCvalue
            sheetOut.cell(row=row1, column=currColA).value = colAvalue + ' (' + model + ')'
        elif colType == 2:  # антенна
            ind = int(sheetOut.cell(row=row1, column=currColE).value.find('AN'))
            model = str(sheetOut.cell(row=row1, column=currColE).value[ind:ind + 9])
            sheetOut.cell(row=row1, column=currColA).value = colAvalue + ' (' + model + ')'
            sheetOut.cell(row=row1, column=currColB).value = 'Antenna'
        else:
            continue
    # удалим лишние столбцы
    sheetOut.delete_cols(6, 2)
    # переместим "ТИП" в конец
    sheetOut.move_range("B1:B2048", rows=0, cols=5)
    sheetOut.delete_cols(2, 1)
    # вставим заглавную строку
    sheetOut.insert_rows(1)
    sheetOut.cell(row=1, column=currColA).value = 'MODEL'
    sheetOut.cell(row=1, column=currColB).value = 'BOM'
    sheetOut.cell(row=1, column=currColC).value = 'NAME'
    sheetOut.cell(row=1, column=currColD).value = 'DESCRIPTION'
    sheetOut.cell(row=1, column=currColE).value = 'CPQ PRICE, USD'
    sheetOut.cell(row=1, column=currColF).value = 'TYPE'
    errors.append('Подготовлен лист продукта -> ' + sheetInName)
    print('Подготовлен лист продукта -> ' + sheetInName)
    return errors


# Определение типа позиции
def find_item_type(strBOM: str, strMODEL: str, strDESCRIPTION: str):
    indTerm1 = strBOM.find('SC')  # проф линейка
    indTerm2 = strBOM.find('CT')  # комм линейка
    indTerm3 = strBOM.find('ZZ')  # ретрансляторы
    indTerm4 = strBOM.find('SP')  # видеорегистраторы
    indAnt1 = strMODEL.find('TQC')  # антенны
    indAnt2 = strMODEL.find('EF')  # антенны
    indAnt3 = strMODEL.find('GP')  # антенны
    indAnt4 = strMODEL.find('AF')  # антенны
    indAnt5 = strMODEL.find('B4')  # антенны
    indAnt6 = strMODEL.find('ODP')  # антенны
    indAnt7 = strMODEL.find('VCD')  # антенны
    indAnt8 = strMODEL.find('CS')  # антенны
    indAnt9 = strMODEL.find('TQJ')  # антенны
    indAnt10 = strMODEL.find('KBT')  # антенны
    indAnt11 = strMODEL.find('DAMA')  # антенны
    indAnt12 = strMODEL.find('CTZ')  # антенны
    indAnt13 = strMODEL.find('M110')  # антенны
    indAnt21 = strDESCRIPTION.find('TQC')  # антенны
    indAnt22 = strDESCRIPTION.find('SC2')  # антенны
    indAnt23 = strDESCRIPTION.find('SC3')  # антенны
    indAnt24 = strDESCRIPTION.find('LC4')  # антенны

    if (indTerm1 != -1) or (indTerm2 != -1) or (indTerm3 != -1) or (indTerm4 != -1):
        return 1  # терминалы
    elif ((indAnt1 != -1) or (indAnt2 != -1) or (indAnt3 != -1) or (indAnt4 != -1) or (indAnt5 != -1)
          or (indAnt6 != -1) or (indAnt7 != -1) or (indAnt8 != -1) or (indAnt9 != -1) or (indAnt10 != -1)
          or (indAnt11 != -1) or (indAnt12 != -1) or (indAnt13 != -1) or (indAnt21 != -1) or (indAnt22 != -1)) \
            or (indAnt23 != -1) or (indAnt24 != -1):
        return 2  # антенны
    else:
        return 0  # аксессуары или лицензии


def copy_cell(source_cell, tgtCoord, tgtSheet):
    tgtSheet[tgtCoord].value = source_cell.value
    # if source_cell.has_style:
    #     tgtSheet[tgtCoord]._style = copy(source_cell._style)
    return tgtSheet[tgtCoord]


def get_cellLetter(idx_row: int, idx_column: int):
    cell_letter = get_column_letter((idx_column)) + str(idx_row)
    return cell_letter


def copy_row(sheet_input, idx_input_row, sheet_output, idx_output_row):
    input_row = sheet_input[idx_input_row]
    for idxColIn, cellIn in enumerate(input_row, start=1):
        # cell_letter_input = get_cellLetter(idx_input_row, idxColIn)
        cell_letter_output = get_cellLetter(idx_output_row, idxColIn)
        row = copy_cell(cellIn, cell_letter_output, sheet_output)


########################################################################################################################
# Все листы на один лист и сделать девайсы уникальными
def onePagePrice(priceCombine_filename: str, priceOnepage_filename: str, isOnlyAntenna: int):
    wbIn = openpyxl.load_workbook(priceCombine_filename, data_only=True)
    wbOut = openpyxl.Workbook()
    sheetOut = wbOut.active
    numberStyle = NamedStyle(name='numberStyle')
    numberStyle.number_format = '0.00'
    numberStyle.font = Font(name='Consolas')
    idxRowOut = 1
    for n, nameIn in enumerate(wbIn.sheetnames):
        sheetIn = wbIn[nameIn]
        for idxRowIn, rowIn in enumerate(sheetIn, start=1):
            if isOnlyAntenna == 0:
                copy_row(sheetIn, idxRowIn, sheetOut, idxRowOut)
                idxRowOut = idxRowOut + 1
            else:
                rowTYPE = str(rowIn[7].value)
                if rowTYPE == 'Antenna':
                    copy_row(sheetIn, idxRowIn, sheetOut, idxRowOut)
                    idxRowOut = idxRowOut + 1
                    # rowIn[1].style = numberStyle
                else:
                    continue
    wbIn.close()
    # вставим заглавную строку
    sheetOut.insert_rows(1)
    sheetOut.cell(row=1, column=1).value = 'BOM'
    sheetOut.cell(row=1, column=2).value = 'MODEL RUS'
    sheetOut.cell(row=1, column=3).value = 'DESCRIPTION ENG'
    sheetOut.cell(row=1, column=4).value = 'DESCRIPTION RUS'
    sheetOut.cell(row=1, column=5).value = 'GPL, CNY'
    sheetOut.cell(row=1, column=6).value = 'РРЦ, CNY'
    sheetOut.cell(row=1, column=7).value = 'COMMENTS'
    sheetOut.cell(row=1, column=8).value = 'MAINTYPE'
    sheetOut.cell(row=1, column=9).value = 'SUBTYPE'
    sheetOut.cell(row=1, column=10).value = 'FREQUENCY BAND'
    wbOut.save(priceOnepage_filename)
    wbOut.close()
    dfPrices = pd.read_excel(priceOnepage_filename, header=0, sheet_name='Sheet', dtype={'BOM': str})
    dfPrices = dfPrices.sort_values(by="DESCRIPTION RUS")
    dfPrices.drop_duplicates(inplace=True)
    dfPrices.to_excel(priceOnepage_filename, index=False)
    print('Успех! ' + onePagePrice.__name__)

#######################################################################################################################
# Конвертов их XLSX в XLS
def convertXLSX2XLS(my_xlsx_excel_file):
    from win32com.client import Dispatch
    my_xlsx_excel_file = my_xlsx_excel_file.replace('/', '\\\\')
    xl = Dispatch('Excel.Application')
    wb = xl.Workbooks.Add(my_xlsx_excel_file)
    newName = my_xlsx_excel_file[:-1]
    wb.SaveAs(newName, FileFormat=56)
    xl.Quit()
    print(newName)
    return newName


#######################################################################################################################
# Русификация прайсов
def setRusPrices(HYTES_PriceList_Filename: str, cpqPricesCombine_Filename: str):
    wbCOMBINE = openpyxl.load_workbook(cpqPricesCombine_Filename, data_only=True)
    mySheets = wbCOMBINE.sheetnames
    wbCOMBINE.close()
    idxXLSX = cpqPricesCombine_Filename.find('.xlsx')-15
    rusPrice_Filename = cpqPricesCombine_Filename[:idxXLSX] + '-RUS' + cpqPricesCombine_Filename[idxXLSX:]
    wbRUS = openpyxl.Workbook()
    wbRUS.save(filename=rusPrice_Filename)
    wbRUS.close()
    dfHYTES = pd.read_excel(HYTES_PriceList_Filename, sheet_name='Sheet1')
    for n, item1 in enumerate(mySheets):
        dfCOMBINE = pd.read_excel(cpqPricesCombine_Filename, sheet_name=item1, dtype={'BOM': str})
        joined_df = pd.merge(dfCOMBINE, dfHYTES, on='BOM', how='left')
        joined_df = joined_df.drop(
            columns=['MODEL_x', 'NAME_x', 'DESCRIPTION_x', 'MODEL_y', 'NAME_y', 'DESCRIPTION_y',
                     'Old CPQ, USD', 'New CPQ, USD', 'TYPE'])
        joined_df = joined_df[['BOM', 'MODEL RUS', 'DESCRIPTION ENG', 'DESCRIPTION RUS',
                               'GPL, CNY', 'РРЦ, CNY', 'COMMENTS', 'MAINTYPE',
                               'SUBTYPE', 'FREQUENCY BAND']]
        joined_df = joined_df.sort_values(by="DESCRIPTION RUS")
        with pd.ExcelWriter(rusPrice_Filename, engine='openpyxl', mode='a') as writer:
            joined_df.to_excel(writer, index=False, sheet_name=item1)
    # удалим пустой лист из документа
    wbRUS = openpyxl.load_workbook(rusPrice_Filename, data_only=True)
    del wbRUS['Sheet']
    wbRUS.save(filename=rusPrice_Filename)
    wbRUS.close()
    checkDiscount(rusPrice_Filename)
    print('Успех! ' + setRusPrices.__name__)
    return rusPrice_Filename

#######################################################################################################################
# Поиск и исправление цен на акционные товары
def checkDiscount(rusPrice_Filename: str):
    errors = []
    wbCHECKED = openpyxl.load_workbook(rusPrice_Filename, data_only=True)
    checkSheets = wbCHECKED.sheetnames
    for n, item1 in enumerate(checkSheets):
        sheetIn = wbCHECKED[item1]
        rowsIn = sheetIn.max_row+1
        for row1 in range(2, rowsIn):
            rowComment = str(sheetIn.cell(row=row1, column=7).value)
            idxDiscount = rowComment.find('акция, ')
            if idxDiscount != -1:
                discountPrice = rowComment[7+idxDiscount:]
                product = str(sheetIn.cell(row=row1, column=1).value)
                errors.append('Акция на ' + product + '. Новая цена = ' + discountPrice)
                print('Акция на ' + product + '. Новая цена = ' + discountPrice)
                sheetIn.cell(row=row1, column=6).value = int(discountPrice)
    wbCHECKED.save(filename=rusPrice_Filename)
    wbCHECKED.close()
    return  errors

#######################################################################################################################
# Поиск новых позиций BOM
def findNewBOM(cpqPricesCombine_Filename: str):
    idxXLSX = cpqPricesCombine_Filename.find('.xlsx')-15
    rusPrice_Filename = cpqPricesCombine_Filename[:idxXLSX] + '-RUS' + cpqPricesCombine_Filename[idxXLSX:]
    newBOM_Filename = cpqPricesCombine_Filename[:idxXLSX] + '-newBOM' + cpqPricesCombine_Filename[idxXLSX:]
    wbRUS = openpyxl.load_workbook(rusPrice_Filename, data_only=True)
    wbNewBOM = openpyxl.Workbook()
    wbNewBOM.save(filename=newBOM_Filename)
    wbNewBOM.close()
    shNames = wbRUS.sheetnames
    df = pd.read_excel(cpqPricesCombine_Filename, sheet_name=shNames[1])
    dfAll = df[df['BOM'] == '1']
    for n, nameIn in enumerate(shNames):
        sheetRUS = wbRUS[nameIn]
        for idxRowIn, rowIn in enumerate(sheetRUS, start=1):
            isNewBOM = str(rowIn[3].value)
            if (isNewBOM == 'None'):
                # print(' cell4=' + str(rowIn[4].value) + ' cell3=' + str(rowIn[3].value) + ' cell2=' + str(rowIn[2].value))
                newBOM = (rowIn[0].value)
                df = pd.read_excel(cpqPricesCombine_Filename, sheet_name=nameIn)
                dfNewBOM = df[df['BOM'].astype(str) == newBOM]
                frames = [dfAll, dfNewBOM]
                dfAll = pd.concat(frames)
            else:
                continue
    wbRUS.close()
    dfAll=dfAll.drop_duplicates().astype(str)
    with pd.ExcelWriter(newBOM_Filename, engine='openpyxl', mode='a') as writer:
        dfAll.to_excel(writer, index=False)
    wbNewBOM = openpyxl.load_workbook(newBOM_Filename, data_only=True)
    del wbNewBOM['Sheet']
    sheet = wbNewBOM.worksheets[0]
    row_count = sheet.max_row-1
    if row_count == 0:
        try:
            os.remove(newBOM_Filename)
        except OSError:
            pass
    else:
        wbNewBOM.save(filename=newBOM_Filename)
        print(newBOM_Filename)
    wbNewBOM.close()
    print('Успех! ' + findNewBOM.__name__)
    return row_count


########################################################################################################################
# Разделение прайса на вкладки
def dividePriceBySheets(Combined_Price_Filename: str):
    wbCOMBINE = openpyxl.load_workbook(Combined_Price_Filename, data_only=True)
    nameIn = wbCOMBINE.sheetnames[0]
    sheetIn = wbCOMBINE[nameIn]
    wbDIVIDE = openpyxl.Workbook()
    idxXLSX = Combined_Price_Filename.find('.xlsx')
    divide_Price_Filename = Combined_Price_Filename[:idxXLSX] + '-DIV' + Combined_Price_Filename[idxXLSX:]
    for idxRowIn, rowIn in enumerate(sheetIn, start=1):
        TYPE = str(rowIn[7].value)
        SUBTYPE = str(rowIn[8].value)
        FREQBAND = str(rowIn[9].value)
        try:
            sheetOut = wbDIVIDE[SUBTYPE]
        except Exception as myError:
            sheetOut = wbDIVIDE.create_sheet(SUBTYPE)
            copy_row(sheetIn, 1, sheetOut, 1)
            # sheetOut[1] = rowTitle
            print(str(myError) + '--> Создаю лист: ' + SUBTYPE)
        idxRowOut = sheetOut.max_row + 1
        copy_row(sheetIn, idxRowIn, sheetOut, idxRowOut)
    wbCOMBINE.close()
    del wbDIVIDE['Sheet']
    del wbDIVIDE['SUBTYPE']
    wbDIVIDE.save(filename=divide_Price_Filename)
    wbDIVIDE.close()
    print('Успех! ' + dividePriceBySheets.__name__)
    return divide_Price_Filename


########################################################################################################################
# Сохранение всех картинок из файла Excel в папку (уже не используется)
def saveImagesExcel2PNG(wb_Filename: str):
    from openpyxl_image_loader import SheetImageLoader
    dir2save = os.path.dirname(os.path.abspath(wb_Filename)) + '/00_Images/'
    if not os.path.exists(dir2save):
        os.makedirs(dir2save)
    wbHYTES = openpyxl.load_workbook(wb_Filename, data_only=True)
    # loading the Excel File and the sheet
    sheets = wbHYTES.sheetnames
    for n, nameSheet1 in enumerate(sheets):
        sheet1 = wbHYTES[nameSheet1]
        image_loader = SheetImageLoader(sheet1)
        # get the image (put the cell you need instead of 'A1')
        rowsIn = sheet1.max_row
        for row1 in range(1, rowsIn):
            cell1 = sheet1.cell(row1, 10).coordinate
            if image_loader.image_in(cell1):
                modelName = str(sheet1.cell(row1, 11).value)
                if modelName == 'None':
                    continue
                imageName = dir2save + modelName + '.PNG'
                try:
                    image1 = image_loader.get(cell1)
                    image1.save(imageName)
                except Exception as e:
                    print(str(e) + ' --> ' + imageName)
            else:
                continue
    print('Успех! ' + saveImagesExcel2PNG.__name__)


#######################################################################################################################
# Преобразование фалов EXCEL в WORD
def processTerminal(i: int, currSheet, symbolGPL, symbolRRC, table_contents_Terminal: list,
                    table_contents_Terminal1: list, table_contents_Terminal3: list, table_contents_Terminal4: list, table_contents_Terminal8: list,
                    table_contents_TerminalEXTRA1: list, table_contents_TerminalEXTRA2: list, table_contents_TerminalEXTRA3: list):
    start = time.time()  ## точка отсчета времени
    currTYPE = str(currSheet.cell(i, 9).value)
    currMODEL = str(currSheet.cell(i, 2).value)
    currFREQBAND = str(currSheet.cell(i, 10).value)
    currBand = currFREQBAND[-1:]
    currComment = str(currSheet.cell(i, 7).value)
    isEXTRA1 = currComment.__contains__('EXTRA1')
    isEXTRA2 = currComment.__contains__('EXTRA2')
    isEXTRA3 = currComment.__contains__('EXTRA3')
    if isEXTRA1 == True:
        table_contents_TerminalEXTRA1.append({
            'Model': currSheet.cell(i, 2).value,
            'DescriptionEng': currSheet.cell(i, 3).value,
            'DescriptionRus': currSheet.cell(i, 4).value,
            'PriceGPL': symbolGPL+str(currSheet.cell(i, 5).value),
            'PriceRRC': symbolRRC+str(currSheet.cell(i, 6).value)
        })
    elif isEXTRA2 == True:
        table_contents_TerminalEXTRA2.append({
            'Model': currSheet.cell(i, 2).value,
            'DescriptionEng': currSheet.cell(i, 3).value,
            'DescriptionRus': currSheet.cell(i, 4).value,
            'PriceGPL': symbolGPL+str(currSheet.cell(i, 5).value),
            'PriceRRC': symbolRRC+str(currSheet.cell(i, 6).value)
        })
    elif isEXTRA3 == True:
        table_contents_TerminalEXTRA3.append({
            'Model': currSheet.cell(i, 2).value,
            'DescriptionEng': currSheet.cell(i, 3).value,
            'DescriptionRus': currSheet.cell(i, 4).value,
            'PriceGPL': symbolGPL+str(currSheet.cell(i, 5).value),
            'PriceRRC': symbolRRC+str(currSheet.cell(i, 6).value)
        })
    else:
        if currBand == '1':
            table_contents_Terminal1.append({
                'Model': currSheet.cell(i, 2).value,
                'DescriptionEng': currSheet.cell(i, 3).value,
                'DescriptionRus': currSheet.cell(i, 4).value,
                'PriceGPL': symbolGPL+str(currSheet.cell(i, 5).value),
                'PriceRRC': symbolRRC+str(currSheet.cell(i, 6).value)
            })
        elif currBand == '3':
            table_contents_Terminal3.append({
                'Model': currSheet.cell(i, 2).value,
                'DescriptionEng': currSheet.cell(i, 3).value,
                'DescriptionRus': currSheet.cell(i, 4).value,
                'PriceGPL': symbolGPL+str(currSheet.cell(i, 5).value),
                'PriceRRC': symbolRRC+str(currSheet.cell(i, 6).value)
            })
        elif currBand == '4':
            table_contents_Terminal4.append({
                'Model': currSheet.cell(i, 2).value,
                'DescriptionEng': currSheet.cell(i, 3).value,
                'DescriptionRus': currSheet.cell(i, 4).value,
                'PriceGPL': symbolGPL+str(currSheet.cell(i, 5).value),
                'PriceRRC': symbolRRC+str(currSheet.cell(i, 6).value)
            })
        elif currBand == '8':
            table_contents_Terminal8.append({
                'Model': currSheet.cell(i, 2).value,
                'DescriptionEng': currSheet.cell(i, 3).value,
                'DescriptionRus': currSheet.cell(i, 4).value,
                'PriceGPL': symbolGPL+str(currSheet.cell(i, 5).value),
                'PriceRRC': symbolRRC+str(currSheet.cell(i, 6).value)
            })
        else:
            table_contents_Terminal.append({
                'Model': currSheet.cell(i, 2).value,
                'DescriptionEng': currSheet.cell(i, 3).value,
                'DescriptionRus': currSheet.cell(i, 4).value,
                'PriceGPL': symbolGPL+str(currSheet.cell(i, 5).value),
                'PriceRRC': symbolRRC+str(currSheet.cell(i, 6).value)
            })
    end = time.time() - start  ## собственно время работы программы
    # print(str(end) + ': processTerminal -> ' + currMODEL)  ## вывод времени

def processSoftware(i: int, currSheet, symbolGPL, symbolRRC, table_contents_Software: list,
                    table_contents_SoftwareEXTRA1: list, table_contents_SoftwareEXTRA2: list, table_contents_SoftwareEXTRA3: list):
    start = time.time()  ## точка отсчета времени
    currMODEL = str(currSheet.cell(i, 2).value)
    currComment = str(currSheet.cell(i, 7).value)
    isEXTRA1 = currComment.__contains__('EXTRA1')
    isEXTRA2 = currComment.__contains__('EXTRA2')
    isEXTRA3 = currComment.__contains__('EXTRA3')
    if isEXTRA1 == True:
        table_contents_SoftwareEXTRA1.append({
            'Model': currSheet.cell(i, 2).value,
            'DescriptionEng': currSheet.cell(i, 3).value,
            'DescriptionRus': currSheet.cell(i, 4).value,
            'PriceGPL': symbolGPL+str(currSheet.cell(i, 5).value),
            'PriceRRC': symbolRRC+str(currSheet.cell(i, 6).value)
        })
    elif isEXTRA2 == True:
        table_contents_SoftwareEXTRA2.append({
            'Model': currSheet.cell(i, 2).value,
            'DescriptionEng': currSheet.cell(i, 3).value,
            'DescriptionRus': currSheet.cell(i, 4).value,
            'PriceGPL': symbolGPL+str(currSheet.cell(i, 5).value),
            'PriceRRC': symbolRRC+str(currSheet.cell(i, 6).value)
        })
    elif isEXTRA3 == True:
        table_contents_SoftwareEXTRA3.append({
            'Model': currSheet.cell(i, 2).value,
            'DescriptionEng': currSheet.cell(i, 3).value,
            'DescriptionRus': currSheet.cell(i, 4).value,
            'PriceGPL': symbolGPL+str(currSheet.cell(i, 5).value),
            'PriceRRC': symbolRRC+str(currSheet.cell(i, 6).value)
        })
    else:
        table_contents_Software.append({
            'Model': currSheet.cell(i, 2).value,
            'DescriptionEng': currSheet.cell(i, 3).value,
            'DescriptionRus': currSheet.cell(i, 4).value,
            'PriceGPL': symbolGPL + str(currSheet.cell(i, 5).value),
            'PriceRRC': symbolRRC + str(currSheet.cell(i, 6).value)
        })
    end = time.time() - start  ## собственно время работы программы
    # print(str(end) + ': processTerminal -> ' + currMODEL)  ## вывод времени

def processAccessory(i: int, currSheet, symbolGPL, symbolRRC, docxGPL, docxRRC, dirImage,
                     table_contents_Accessory_Antenna: list, table_contents_Accessory_Audio: list, table_contents_Accessory_Cable: list,
                     table_contents_Accessory_Case: list, table_contents_Accessory_Mount: list, table_contents_Accessory_PowerSupply: list,
                     table_contents_Accessory_Other: list):
    from docx.shared import Cm
    errors = []
    start = time.time()  ## точка отсчета времени
    currMODEL = str(currSheet.cell(i, 2).value)
    currSUBTYPE = str(currSheet.cell(i, 9).value)
    currImageEmptyGPL = InlineImage(docxGPL, dirImage + '/empty.png', width=Cm(2))
    currImageEmptyRRC = InlineImage(docxRRC, dirImage + '/empty.png', width=Cm(2))
    currImageName = dirImage + '/' + currMODEL + '.PNG'
    if pathlib.Path(currImageName).is_file() == True:
        try:
            currImageGPL = InlineImage(docxGPL, currImageName, width=Cm(2))
            currImageRRC = InlineImage(docxRRC, currImageName, width=Cm(2))
        except:
            currImageGPL = currImageEmptyGPL
            currImageRRC = currImageEmptyRRC
            errors.append('error open image --> ' + currImageName)
            print('error open image --> ' + currImageName)
    else:
        currImageGPL = currImageEmptyGPL
        currImageRRC = currImageEmptyRRC
        errors.append('no such image --> ' + currImageName)
        print('no such image --> ' + currImageName)
    if currSUBTYPE == 'Antenna':
        table_contents_Accessory_Antenna.append({
            'Model': currSheet.cell(i, 2).value,
            'DescriptionEng': currSheet.cell(i, 3).value,
            'DescriptionRus': currSheet.cell(i, 4).value,
            'PriceGPL': symbolGPL + str(currSheet.cell(i, 5).value),
            'PriceRRC': symbolRRC + str(currSheet.cell(i, 6).value),
            'ImageGPL': currImageGPL,
            'ImageRRC': currImageRRC
        })
    elif currSUBTYPE == 'Audio':
        table_contents_Accessory_Audio.append({
            'Model': currSheet.cell(i, 2).value,
            'DescriptionEng': currSheet.cell(i, 3).value,
            'DescriptionRus': currSheet.cell(i, 4).value,
            'PriceGPL': symbolGPL + str(currSheet.cell(i, 5).value),
            'PriceRRC': symbolRRC + str(currSheet.cell(i, 6).value),
            'ImageGPL': currImageGPL,
            'ImageRRC': currImageRRC
        })
    elif currSUBTYPE == 'Cable':
        table_contents_Accessory_Cable.append({
            'Model': currSheet.cell(i, 2).value,
            'DescriptionEng': currSheet.cell(i, 3).value,
            'DescriptionRus': currSheet.cell(i, 4).value,
            'PriceGPL': symbolGPL + str(currSheet.cell(i, 5).value),
            'PriceRRC': symbolRRC + str(currSheet.cell(i, 6).value),
            'ImageGPL': currImageGPL,
            'ImageRRC': currImageRRC
        })
    elif currSUBTYPE == 'Case':
        table_contents_Accessory_Case.append({
            'Model': currSheet.cell(i, 2).value,
            'DescriptionEng': currSheet.cell(i, 3).value,
            'DescriptionRus': currSheet.cell(i, 4).value,
            'PriceGPL': symbolGPL + str(currSheet.cell(i, 5).value),
            'PriceRRC': symbolRRC + str(currSheet.cell(i, 6).value),
            'ImageGPL': currImageGPL,
            'ImageRRC': currImageRRC
        })
    elif currSUBTYPE == 'Mount':
        table_contents_Accessory_Mount.append({
            'Model': currSheet.cell(i, 2).value,
            'DescriptionEng': currSheet.cell(i, 3).value,
            'DescriptionRus': currSheet.cell(i, 4).value,
            'PriceGPL': symbolGPL + str(currSheet.cell(i, 5).value),
            'PriceRRC': symbolRRC + str(currSheet.cell(i, 6).value),
            'ImageGPL': currImageGPL,
            'ImageRRC': currImageRRC
        })
    elif currSUBTYPE == 'PowerSupply':
        table_contents_Accessory_PowerSupply.append({
            'Model': currSheet.cell(i, 2).value,
            'DescriptionEng': currSheet.cell(i, 3).value,
            'DescriptionRus': currSheet.cell(i, 4).value,
            'PriceGPL': symbolGPL + str(currSheet.cell(i, 5).value),
            'PriceRRC': symbolRRC + str(currSheet.cell(i, 6).value),
            'ImageGPL': currImageGPL,
            'ImageRRC': currImageRRC
        })
    else:
        table_contents_Accessory_Other.append({
            'Model': currSheet.cell(i, 2).value,
            'DescriptionEng': currSheet.cell(i, 3).value,
            'DescriptionRus': currSheet.cell(i, 4).value,
            'PriceGPL': symbolGPL + str(currSheet.cell(i, 5).value),
            'PriceRRC': symbolRRC + str(currSheet.cell(i, 6).value),
            'ImageGPL': currImageGPL,
            'ImageRRC': currImageRRC
        })
    # end = time.time() - start  ## собственно время работы программы
    # print(str(end) + ': processAccessory -> ' + currMODEL)  ## вывод времени
    return errors

def excel2word(excel_File_Rus: str, templateGPL: str, templateRRC: str, dirImage: str, price_date: str, price_version: str):
    from docx2pdf import convert
    import comtypes.client
    wdFormatPDF = 17
    errors = []
    startAll = time.time()  ## точка отсчета времени
    # start = time.time()  ## точка отсчета времени
    # end = time.time() - start  ## собственно время работы программы
    # print(str(end))  ## вывод времени
    wbIn = openpyxl.load_workbook(excel_File_Rus, data_only=True)
    docxGPL = DocxTemplate(templateGPL)
    docxRRC = DocxTemplate(templateRRC)
    context = {}
    symbolGPL = "¥"
    symbolRRC = "¥"
    for n, sheetName in enumerate(wbIn.sheetnames):
        currSheet = wbIn[sheetName]
        table_contents_Terminal = []
        table_contents_Terminal1 = []
        table_contents_Terminal3 = []
        table_contents_Terminal4 = []
        table_contents_Terminal8 = []
        table_contents_TerminalEXTRA1 = []
        table_contents_TerminalEXTRA2 = []
        table_contents_TerminalEXTRA3 = []
        table_contents_Software = []
        table_contents_SoftwareEXTRA1 = []
        table_contents_SoftwareEXTRA2 = []
        table_contents_SoftwareEXTRA3 = []
        table_contents_Accessory_Antenna = []
        table_contents_Accessory_Audio = []
        table_contents_Accessory_Cable = []
        table_contents_Accessory_Case = []
        table_contents_Accessory_Mount = []
        table_contents_Accessory_Other = []
        table_contents_Accessory_PowerSupply = []
        table_contents_Antenna = []
        for i in range(2, currSheet.max_row + 1):
            currTYPE = str(currSheet.cell(i, 8).value)
            currMODEL = str(currSheet.cell(i, 2).value)
            indErr1 = currMODEL.find('not certified')  # ошибка
            indErr2 = currMODEL.find('None')  # ошибка
            if (indErr1 != -1) or (indErr2 != -1):
                continue
            if currTYPE == 'Terminal':
                processTerminal(i, currSheet, symbolGPL, symbolRRC,
                                table_contents_Terminal, table_contents_Terminal1, table_contents_Terminal3, table_contents_Terminal4, table_contents_Terminal8,
                                table_contents_TerminalEXTRA1, table_contents_TerminalEXTRA2, table_contents_TerminalEXTRA3)
            elif currTYPE == 'Software':
                processSoftware(i, currSheet, symbolGPL, symbolRRC,
                                table_contents_Software, table_contents_SoftwareEXTRA1, table_contents_SoftwareEXTRA2, table_contents_SoftwareEXTRA3)
            elif currTYPE == 'Accessory':
                errors.append(processAccessory(i, currSheet, symbolGPL, symbolRRC, docxGPL, docxRRC, dirImage,
                table_contents_Accessory_Antenna, table_contents_Accessory_Audio, table_contents_Accessory_Cable,
                table_contents_Accessory_Case, table_contents_Accessory_Mount, table_contents_Accessory_PowerSupply,
                table_contents_Accessory_Other))
            elif (currTYPE == 'Antenna') or (currTYPE == 'AntennaMobile'):
                table_contents_Antenna.append({
                    'Model': currSheet.cell(i, 2).value,
                    'DescriptionEng': currSheet.cell(i, 3).value,
                    'DescriptionRus': currSheet.cell(i, 4).value,
                    'PriceGPL': symbolGPL + str(currSheet.cell(i, 5).value),
                    'PriceRRC': symbolRRC + str(currSheet.cell(i, 6).value)
                })
            else:
                errors.append('неизвестный тип = ' + currTYPE)
                print('неизвестный тип = ' + currTYPE)
        curTableContentsTerminal = 'table_contents_' + sheetName
        curTableContentsTerminal1 = 'table_contents_VHF_' + sheetName
        curTableContentsTerminal3 = 'table_contents_300_' + sheetName
        curTableContentsTerminal4 = 'table_contents_UHF_' + sheetName
        curTableContentsTerminal8 = 'table_contents_800_' + sheetName
        curTableContentsTerminalEXTRA1 = 'table_contents_EXTRA1_' + sheetName
        curTableContentsTerminalEXTRA2 = 'table_contents_EXTRA2_' + sheetName
        curTableContentsTerminalEXTRA3 = 'table_contents_EXTRA3_' + sheetName
        curTableContentsLic = 'table_contents_lic_' + sheetName
        curTableContentsLicEXTRA1 = 'table_contents_lic_EXTRA1_' + sheetName
        curTableContentsLicEXTRA2 = 'table_contents_lic_EXTRA2_' + sheetName
        curTableContentsLicEXTRA3 = 'table_contents_lic_EXTRA3_' + sheetName
        curTableContentsAnt = 'table_contents_ant_' + sheetName
        curTableContentsAccAnt = 'table_contents_acc_ant_' + sheetName
        curTableContentsAccAud = 'table_contents_acc_aud_' + sheetName
        curTableContentsAccCab = 'table_contents_acc_cab_' + sheetName
        curTableContentsAccCas = 'table_contents_acc_cas_' + sheetName
        curTableContentsAccMou = 'table_contents_acc_mou_' + sheetName
        curTableContentsAccPow = 'table_contents_acc_pow_' + sheetName
        curTableContentsAccOth = 'table_contents_acc_oth_' + sheetName
        context1 = {
            curTableContentsTerminal: table_contents_Terminal,
            curTableContentsTerminal1: table_contents_Terminal1,
            curTableContentsTerminal3: table_contents_Terminal3,
            curTableContentsTerminal4: table_contents_Terminal4,
            curTableContentsTerminal8: table_contents_Terminal8,
            curTableContentsTerminalEXTRA1: table_contents_TerminalEXTRA1,
            curTableContentsTerminalEXTRA2: table_contents_TerminalEXTRA2,
            curTableContentsTerminalEXTRA3: table_contents_TerminalEXTRA3,
            curTableContentsLic: table_contents_Software,
            curTableContentsLicEXTRA1: table_contents_SoftwareEXTRA1,
            curTableContentsLicEXTRA2: table_contents_SoftwareEXTRA2,
            curTableContentsLicEXTRA3: table_contents_SoftwareEXTRA3,
            curTableContentsAnt: table_contents_Antenna,
            curTableContentsAccAnt: table_contents_Accessory_Antenna,
            curTableContentsAccAud: table_contents_Accessory_Audio,
            curTableContentsAccCab: table_contents_Accessory_Cable,
            curTableContentsAccCas: table_contents_Accessory_Case,
            curTableContentsAccMou: table_contents_Accessory_Mount,
            curTableContentsAccPow: table_contents_Accessory_PowerSupply,
            curTableContentsAccOth: table_contents_Accessory_Other
        }
        context.update(context1)
        start = time.time()  ## точка отсчета времени
        docxGPL.render(context)
        end = time.time() - start  ## собственно время работы программы
        errors.append('Успех! Генерация файла ' + sheetName + ' docx-GPL завершена! Время генерации: '+str(end) + ',c')
        print('Успех! Генерация файла ' + sheetName + ' docx-GPL завершена! Время генерации: '+str(end) + ',c')
        start = time.time()  ## точка отсчета времени
        docxRRC.render(context)
        end = time.time() - start  ## собственно время работы программы
        errors.append('Успех! Генерация файла ' + sheetName + ' docx-РРЦ завершена! Время генерации: '+str(end) + ',c')
        print('Успех! Генерация файла ' + sheetName + ' docx-РРЦ завершена! Время генерации: '+str(end) + ',c')
    wbIn.close()
    filenames = makeStandardDocxName(excel_File_Rus, price_date, price_version)
    errors.append('Сохраняю в DOCX.......')
    print('Сохраняю в DOCX.......')
    docxGPL.save(filenames[0])
    docxRRC.save(filenames[1])
    update_toc(filenames[0])
    update_toc(filenames[1])
    errors.append('Конвертирую в PDF.......')
    print('Конвертирую в PDF.......')
    convert(filenames[0])
    convert(filenames[1])
    print('Успех! ' + excel2word.__name__)
    endAll = (time.time() - startAll)/60  ## собственно время работы программы
    print('Время конвертации: '+str(endAll) + ',мин')
    errors.append('Время конвертации: '+str(endAll) + ',мин')
    return errors

def makeStandardDocxName(old_filename: str, price_date: str, price_version: str):
    dir2save = os.path.dirname(os.path.abspath(old_filename)) + '\\'
    filenameOnly = os.path.basename(os.path.abspath(old_filename))
    idxTYPE = filenameOnly.find('-RUS')
    fileType = filenameOnly[:idxTYPE]
    filename_GPL = dir2save + str(price_date) + ' Price-book Hytera ' + fileType + ' (GPL, без НДС, для партнеров) ' + str(price_version) + '.docx'
    filename_RRC = dir2save + str(price_date) + ' Price-book Hytera ' + fileType + ' (РРЦ, с НДС, для заказчиков) ' + str(price_version) + '.docx'
    filenames = [filename_GPL, filename_RRC]
    # print(filenames[0])
    # print(filenames[1])
    return filenames

########################################################################################################################
# Сделать из шаблона GPL шаблон РРЦ
def convertGPL2RRC(word_template: str, version: str, dated:str):
    import docx
    word_files = []
    idxXLSX = word_template.find('.docx')
    word_FileGPL = word_template[:idxXLSX] + ' GPL.docx'
    word_FileRRC = word_template[:idxXLSX] + ' РРЦ.docx'
    word_files.append(word_FileGPL)
    word_files.append(word_FileRRC)

    doc = docx.Document(word_template)
    docx_find_replace_text(doc, '[type]', 'GPL, без НДС')
    docx_find_replace_text(doc, '[version]', version)
    docx_find_replace_text(doc, '[dated]', dated)
    doc.save(word_FileGPL)
    docx_find_replace_text(doc, 'GPL, без НДС', 'РРЦ, с НДС')
    docx_find_replace_text(doc, 'ImageGPL', 'ImageRRC')
    docx_find_replace_text(doc, 'PriceGPL', 'PriceRRC')
    doc.save(word_FileRRC)
    return word_files

def update_toc(docx_file):
    # pass
    import win32com.client
    word = win32com.client.DispatchEx("Word.Application")
    doc = word.Documents.Open(docx_file)
    doc.TablesOfContents(1).Update()
    # doc.TablesOfContents(1).UpdatePageNumbers()
    doc.Close(SaveChanges=True)
    # word.Quit()


# функция поиска и замены
def docx_find_replace_text(doc, search_text, replace_text):
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        if paragraph.text.find(search_text) >= 0:
            paragraph.text = paragraph.text.replace(search_text, replace_text)

    paragraphs = list(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraphs.append(paragraph)
    for p in paragraphs:
        if search_text in p.text:
            inline = p.runs
            # Replace strings and retain the same style.
            # The text to be replaced can be split over several runs so
            # search through, identify which runs need to have text replaced
            # then replace the text in those identified
            started = False
            search_index = 0
            # found_runs is a list of (inline index, index of match, length of match)
            found_runs = list()
            found_all = False
            replace_done = False
            for i in range(len(inline)):

                # case 1: found in single run so short circuit the replace
                if search_text in inline[i].text and not started:
                    found_runs.append((i, inline[i].text.find(search_text), len(search_text)))
                    text = inline[i].text.replace(search_text, str(replace_text))
                    inline[i].text = text
                    replace_done = True
                    found_all = True
                    break

                if search_text[search_index] not in inline[i].text and not started:
                    # keep looking ...
                    continue

                # case 2: search for partial text, find first run
                if search_text[search_index] in inline[i].text and inline[i].text[-1] in search_text and not started:
                    # check sequence
                    start_index = inline[i].text.find(search_text[search_index])
                    check_length = len(inline[i].text)
                    for text_index in range(start_index, check_length):
                        if inline[i].text[text_index] != search_text[search_index]:
                            # no match so must be false positive
                            break
                    if search_index == 0:
                        started = True
                    chars_found = check_length - start_index
                    search_index += chars_found
                    found_runs.append((i, start_index, chars_found))
                    if search_index != len(search_text):
                        continue
                    else:
                        # found all chars in search_text
                        found_all = True
                        break

                # case 2: search for partial text, find subsequent run
                if search_text[search_index] in inline[i].text and started and not found_all:
                    # check sequence
                    chars_found = 0
                    check_length = len(inline[i].text)
                    for text_index in range(0, check_length):
                        if inline[i].text[text_index] == search_text[search_index]:
                            search_index += 1
                            chars_found += 1
                        else:
                            break
                    # no match so must be end
                    found_runs.append((i, 0, chars_found))
                    if search_index == len(search_text):
                        found_all = True
                        break

            if found_all and not replace_done:
                for i, item in enumerate(found_runs):
                    index, start, length = [t for t in item]
                    if i == 0:
                        text = inline[index].text.replace(inline[index].text[start:start + length], str(replace_text))
                        inline[index].text = text
                    else:
                        text = inline[index].text.replace(inline[index].text[start:start + length], '')
                        inline[index].text = text
            # print(p.text)

########################################################################################################################
# прочитать датафрейм с картинками
def load_dataframe(dataframe_file_path: str, dataframe_sheet_name: str) -> pd.DataFrame:
    from io import BytesIO
    import base64
    from openpyxl_image_loader import SheetImageLoader
    # By default, it appears that pandas does not read images, as it uses only openpyxl to read
    # the file.  As a result we need to load into memory the dataframe and explicitly load in
    # the images, and then convert all of this to HTML and put it back into the normal
    # dataframe, ready for use.
    pxl_doc = openpyxl.load_workbook(dataframe_file_path)
    pxl_sheet = pxl_doc[dataframe_sheet_name]
    pxl_image_loader = SheetImageLoader(pxl_sheet)
    pd_df = pd.read_excel(dataframe_file_path, sheet_name=dataframe_sheet_name)
    for pd_row_idx, pd_row_data in pd_df.iterrows():
        for pd_column_idx, _pd_cell_data in enumerate(pd_row_data):
            # Offset as openpyxl sheets index by one, and also offset the row index by one more to account for the
            # header row
            pxl_cell_coord_str = pxl_sheet.cell(pd_row_idx + 2, pd_column_idx + 1).coordinate
            if pxl_image_loader.image_in(pxl_cell_coord_str):
                # Now that we have a cell that contains an image, we want to convert it to
                # base64, and it make it nice and HTML, so that it loads in a front end
                pxl_pil_img = pxl_image_loader.get(pxl_cell_coord_str)
                with BytesIO() as pxl_pil_buffered:
                    pxl_pil_img.save(pxl_pil_buffered, format="PNG")
                    # pd_df.iat[pd_row_idx, pd_column_idx] = pxl_pil_img
                    pd_df.iat[pd_row_idx, pd_column_idx] = base64.b64encode(pxl_pil_buffered.getvalue()).decode('utf-8')
                    # pd_df.iat[pd_row_idx, pd_column_idx] = '<img src="data:image/png;base64,' + \
                    #                                            pxl_pil_img_b64_str.decode('utf-8') + \
                    #                                            f'" alt="{pxl_cell_coord_str}" />'
    pxl_doc.close()
    return pd_df
